import re
from docx import Document

def combine_documents_with_regex(devanagari_path, roman_path, output_path):
    # Load the documents
    devanagari_doc = Document(devanagari_path)
    roman_doc = Document(roman_path)
    
    # Initialize the output document
    combined_doc = Document()

    # Regular expression to match verse endings
    verse_end_pattern = re.compile(r'॥ \d{1,2}\.\d{1,2}$')

    # Function to merge paragraphs based on the verse ending pattern
    def merge_paragraphs(doc):
        merged_paragraphs = []
        current_paragraph = ""
        for paragraph in doc.paragraphs:
            current_paragraph += paragraph.text + "\n"
            if verse_end_pattern.search(paragraph.text):
                merged_paragraphs.append(current_paragraph.strip())
                current_paragraph = ""
        return merged_paragraphs

    # Merge paragraphs for each document
    devanagari_paragraphs = merge_paragraphs(devanagari_doc)
    roman_paragraphs = merge_paragraphs(roman_doc)

    # Combine paragraphs into the output document
    for devanagari_para, roman_para in zip(devanagari_paragraphs, roman_paragraphs):
        combined_doc.add_paragraph(devanagari_para)
        para = combined_doc.add_paragraph(roman_para)
        for run in para.runs:
            run.italic = True
        combined_doc.add_page_break()

    # Save the output document
    combined_doc.save(output_path)

# Specify the file paths
devanagari_path = 'Bhagavad-gita in Devanagari Script.docx'
roman_path = 'Bhagavad-gita in Roman Script.docx'
output_path = 'Combined_Bhagavad-Gita_with_Regex.docx'

# Combine the documents using regex to identify paragraphs
combine_documents_with_regex(devanagari_path, roman_path, output_path)


